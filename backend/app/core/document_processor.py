"""
Document Processor - Handles file parsing and text chunking.

Supports: PDF, TXT, DOCX
Uses recursive character-based chunking with overlap.
"""

import os
import re
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Parses documents and splits into overlapping chunks
    suitable for embedding and retrieval.
    """

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def process_file(self, file_path: str) -> list[dict]:
        """Load, extract text, and chunk a document."""
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            text = self._extract_pdf(file_path)
        elif ext == ".txt":
            text = self._extract_txt(file_path)
        elif ext == ".docx":
            text = self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        if not text or not text.strip():
            raise ValueError(f"No text could be extracted from {file_path}")

        # Clean and chunk
        text = self._clean_text(text)
        chunks = self._chunk_text(text)

        return [
            {
                "text": chunk,
                "chunk_index": i,
                "source": path.name,
            }
            for i, chunk in enumerate(chunks)
        ]

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF (fitz) or pdfplumber."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            return "\n".join(text_parts)
        except ImportError:
            pass

        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                return "\n".join(
                    page.extract_text() or "" for page in pdf.pages
                )
        except ImportError:
            pass

        raise ImportError("Install pymupdf or pdfplumber: pip install pymupdf pdfplumber")

    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Could not decode {file_path}")

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            return "\n".join(paragraphs)
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)
        # Remove page headers/footers patterns
        text = re.sub(r"Page \d+ of \d+", "", text, flags=re.IGNORECASE)
        return text.strip()

    def _chunk_text(self, text: str) -> list[str]:
        """
        Recursive character-based chunking with overlap.
        Tries to split on paragraph → sentence → word boundaries.
        """
        separators = ["\n\n", "\n", ". ", "? ", "! ", ", ", " "]
        chunks = self._split_recursive(text, separators)
        return [c.strip() for c in chunks if c.strip()]

    def _split_recursive(self, text: str, separators: list[str]) -> list[str]:
        """Recursively split text using a list of separators."""
        if len(text) <= self.chunk_size:
            return [text]

        separator = separators[0] if separators else " "
        remaining_seps = separators[1:]

        splits = text.split(separator)
        chunks = []
        current = ""

        for split in splits:
            test = current + (separator if current else "") + split
            if len(test) <= self.chunk_size:
                current = test
            else:
                if current:
                    chunks.append(current)
                    # Add overlap from end of previous chunk
                    overlap_text = current[-self.chunk_overlap:] if len(current) > self.chunk_overlap else current
                    current = overlap_text + (separator if overlap_text else "") + split
                else:
                    # Single piece too large, recurse with next separator
                    if remaining_seps:
                        sub_chunks = self._split_recursive(split, remaining_seps)
                        chunks.extend(sub_chunks[:-1])
                        current = sub_chunks[-1] if sub_chunks else ""
                    else:
                        # Hard split
                        while len(split) > self.chunk_size:
                            chunks.append(split[:self.chunk_size])
                            split = split[self.chunk_size - self.chunk_overlap:]
                        current = split

        if current:
            chunks.append(current)

        return chunks
